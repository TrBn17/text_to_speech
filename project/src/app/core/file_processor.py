import io
import base64
from typing import Dict, Any, List, Optional
from pathlib import Path
import tempfile
import os

# Optional imports for file processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class FileProcessor:
    """Service để xử lý các loại file input"""

    def __init__(self):
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_docx,
            'txt': self._process_text,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'bmp': self._process_image,
            'tiff': self._process_image
        }

    async def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Xử lý file và extract text content"""
        try:
            # Get file extension
            file_ext = Path(filename).suffix.lower().lstrip('.')

            if file_ext not in self.supported_extensions:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_ext}",
                    "supported_types": list(self.supported_extensions.keys())
                }

            # Process file based on type
            processor = self.supported_extensions[file_ext]
            result = await processor(file_content, filename)

            return {
                "success": True,
                "filename": filename,
                "file_type": file_ext,
                "content": result.get("content", ""),
                "metadata": result.get("metadata", {}),
                "page_count": result.get("page_count", 1)
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing file {filename}: {str(e)}",
                "filename": filename
            }

    async def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return {
                "content": "",
                "error": "PDF processing libraries not installed. Install with: pip install PyPDF2 pdfplumber"
            }

        try:
            text_content = ""
            metadata = {}

            # Try with pdfplumber first (better text extraction)
            try:
                with io.BytesIO(file_content) as pdf_buffer:
                    import pdfplumber
                    with pdfplumber.open(pdf_buffer) as pdf:
                        pages = []
                        for page_num, page in enumerate(pdf.pages, 1):
                            page_text = page.extract_text() or ""
                            if page_text.strip():
                                pages.append(f"[Page {page_num}]\n{page_text}")

                        text_content = "\n\n".join(pages)
                        metadata = {
                            "page_count": len(pdf.pages),
                            "extraction_method": "pdfplumber"
                        }
            except Exception:
                # Fallback to PyPDF2
                with io.BytesIO(file_content) as pdf_buffer:
                    pdf_reader = PyPDF2.PdfReader(pdf_buffer)
                    pages = []
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            pages.append(f"[Page {page_num}]\n{page_text}")

                    text_content = "\n\n".join(pages)
                    metadata = {
                        "page_count": len(pdf_reader.pages),
                        "extraction_method": "PyPDF2"
                    }

            return {
                "content": text_content,
                "metadata": metadata,
                "page_count": metadata.get("page_count", 0)
            }

        except Exception as e:
            return {
                "content": "",
                "error": f"PDF processing failed: {str(e)}"
            }

    async def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return {
                "content": "",
                "error": "DOCX processing library not installed. Install with: pip install python-docx"
            }

        try:
            with io.BytesIO(file_content) as docx_buffer:
                doc = Document(docx_buffer)

                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            tables_text.append(" | ".join(row_text))

                content = "\n\n".join(paragraphs)
                if tables_text:
                    content += "\n\n[Tables]\n" + "\n".join(tables_text)

                return {
                    "content": content,
                    "metadata": {
                        "paragraphs_count": len(paragraphs),
                        "tables_count": len(doc.tables)
                    }
                }

        except Exception as e:
            return {
                "content": "",
                "error": f"DOCX processing failed: {str(e)}"
            }

    async def _process_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    content = file_content.decode(encoding)
                    return {
                        "content": content,
                        "metadata": {
                            "encoding": encoding,
                            "lines_count": len(content.splitlines())
                        }
                    }
                except UnicodeDecodeError:
                    continue

            return {
                "content": "",
                "error": "Could not decode text file with supported encodings"
            }

        except Exception as e:
            return {
                "content": "",
                "error": f"Text processing failed: {str(e)}"
            }

    async def _process_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            return {
                "content": "",
                "error": "OCR libraries not installed. Install with: pip install Pillow pytesseract"
            }

        try:
            # Save to temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name

            try:
                # Open image and perform OCR
                image = Image.open(tmp_file_path)

                # Convert to RGB if necessary
                if image.mode in ('RGBA', 'LA', 'P'):
                    image = image.convert('RGB')

                # Extract text using OCR
                extracted_text = pytesseract.image_to_string(image)

                return {
                    "content": extracted_text,
                    "metadata": {
                        "image_size": image.size,
                        "image_mode": image.mode,
                        "extraction_method": "pytesseract"
                    }
                }

            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)

        except Exception as e:
            return {
                "content": "",
                "error": f"Image OCR processing failed: {str(e)}"
            }

    def get_supported_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.supported_extensions.keys())

# Global instance
file_processor = FileProcessor()